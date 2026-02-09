from typing import List, Any
from pydantic import BaseModel
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
import shutil
import os
import uuid
from datetime import datetime
import asyncio
import aiofiles
from functools import partial
import json

from app.api import deps
from app.models.user import User
from app.models.document import Document, Chunk
from app.services.vector_store import vector_store
from app.services.ingestion import ingestion_service
from app.services.metadata_extraction import metadata_service
from app.services.retrieval_service import retrieval_service
from app.db.session import AsyncSessionLocal

# Wrapper to run in background with fresh session
async def run_metadata_extraction(memory_id: int, user_id: int, doc_type: str = "memory"):
    async with AsyncSessionLocal() as db:
        try:
            await metadata_service.process_memory_metadata(memory_id, user_id, db, doc_type)
        except Exception as e:
            print(f"Error in background metadata extraction: {e}")

# Text Extraction Libraries
import pdfplumber
import docx
from bs4 import BeautifulSoup

router = APIRouter()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text_from_file(file_path: str, file_type: str) -> str:
    text = ""
    try:
        if file_type == "pdf":
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
        elif file_type == "docx":
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_type in ["txt", "md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif file_type == "html":
             with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                text = soup.get_text()
    except Exception as e:
        print(f"Error extracting text: {e}")
    return text

    return text

@router.post("/upload", response_model=Any)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    """
    Upload a document, extract text, chunk it, and store in Vector DB.
    """
    file_ext = file.filename.split(".")[-1].lower()
    if file_ext not in ["pdf", "docx", "txt", "md", "html"]:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    # Save file temporarily (Async)
    file_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}.{file_ext}")
    try:
        async with aiofiles.open(file_path, 'wb') as out_file:
            while content := await file.read(1024 * 1024):  # Read in 1MB chunks
                await out_file.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save upload: {e}")
        
    # Extract Text (Run in Executor to avoid blocking)
    loop = asyncio.get_event_loop()
    try:
        text = await loop.run_in_executor(None, extract_text_from_file, file_path, file_ext)
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        print(f"Extraction Error: {e}")
        raise HTTPException(status_code=400, detail="Could not extract text from file")
        
    if not text.strip():
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=400, detail="Could not extract text from file (empty)")
        
    # Cleanup file immediately after extraction? 
    # Or keep it? The original code didn't remove it explicitly unless error/cleanup.
    # We should probably keep it if we want to support download later, but for now we remove to save space?
    # Original code didn't show removal logic except in error block or verify script.
    # Let's clean it up to avoid clutter since we store text.
    try:
         os.remove(file_path)
    except:
         pass
        
    # Create Document Record
    document = Document(
        title=file.filename,
        content=text,  # Store extracted text
        source=file.filename,
        file_type=file_ext,
        doc_type="file",  # Mark as file upload
        user_id=current_user.id
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)
    
    # Offload Ingestion to Background Task
    try:
        from app.worker import ingest_memory_task
        print(f"Triggering background ingestion for document {document.id}")
        ingest_memory_task.delay(
            memory_id=document.id,
            user_id=current_user.id,
            content=text,
            title=document.title,
            tags=[],
            source=document.source, # filename
            doc_type="file"
        )
    except Exception as e:
        print(f"Failed to trigger background task: {e}")
        # Fallback to foreground if celery fails is optional, but for now we log.
        
    return {"status": "success", "document_id": document.id, "message": "Document queued for processing"}

from app.services.youtube_service import youtube_service

class YouTubeUpload(BaseModel):
    url: str
    tags: List[str] = []

@router.post("/upload-youtube", response_model=Any)
async def upload_youtube(
    request: YouTubeUpload,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    """
    Ingest a YouTube video transcript.
    """
    # Try to extract transcript
    try:
        transcript = youtube_service.extract_transcript(request.url)
    except Exception as e:
        print(f"Warning: Transcript extraction threw error: {e}")
        transcript = None
        
    # Fetch real title and description
    video_id = youtube_service.get_video_id(request.url)
    title = youtube_service.get_video_title(request.url)
    description = youtube_service.get_video_description(request.url)
    
    final_content = ""
    # Merge Default Tags with Request Tags
    tags = ["youtube"]
    if request.tags:
        tags.extend(request.tags)
        # Dedupe
        tags = list(set(tags))
    
    if transcript and transcript.strip():
        # Happy path
        final_content = transcript
    else:
        # Fallback path
        print(f"Fallback: Using description for {video_id}")
        final_content = f"[Transcript Unavailable - Metadata Only]\n\nDescription:\n{description}"
        tags.append("no-transcript")
        
    # Create Memory Record (showing in Inbox)
    from app.models.memory import Memory
    
    # Logic: Only show in inbox if from extension OR if auto-approve is off (pending)
    # Default Assume Approved/Direct for App, unless implicit setting.
    # User Request: "Only anything coming from extension should end in the inbox"
    # So if "extension" in tags -> show_in_inbox = True.
    # Else -> show_in_inbox = False (direct save).
    
    show_in_inbox = False
    if "extension" in tags:
        show_in_inbox = True
    
    memory = Memory(
        title=title,
        content=final_content,
        source_llm="youtube", 
        user_id=current_user.id,
        tags=tags,
        embedding_id=str(uuid.uuid4()),
        status="pending" if show_in_inbox else "approved", # If skipping inbox, it's approved
        show_in_inbox=show_in_inbox
    )
    db.add(memory)
    await db.commit()
    await db.refresh(memory)
    
    # Offload Ingestion to Background Task
    try:
        from app.worker import ingest_memory_task, process_memory_metadata_task
        print(f"Triggering background ingestion for YouTube Memory {memory.id}")
        
        # Metadata analysis
        process_memory_metadata_task.delay(memory.id, current_user.id)
        
        # Ingest content
        ingest_memory_task.delay(
            memory_id=memory.id,
            user_id=current_user.id,
            content=final_content,
            title=memory.title,
            tags=tags,
            source=request.url,
            doc_type="memory" # Important: treat as memory now
        )
    except Exception as e:
        print(f"Failed to trigger background task: {e}")

    return {"status": "success", "document_id": memory.id, "message": "Video queued."}

@router.get("/", response_model=Any)
async def get_documents(
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    result = await db.execute(select(Document).where(Document.user_id == current_user.id))
    documents = result.scalars().all()
    
    # Convert to dictionaries for serialization
    result_list = []
    for doc in documents:
        result_list.append({
            "id": doc.id,
            "title": doc.title,
            "content": doc.content,
            "source": doc.source,
            "file_type": doc.file_type,
            "doc_type": doc.doc_type,
            "created_at": doc.created_at,
            "updated_at": doc.updated_at,
            "tags": doc.tags
        })
    
    return result_list

@router.delete("/{doc_id}", response_model=Any)
async def delete_document(
    doc_id: int,
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == current_user.id))
    document = result.scalars().first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
        
    # Delete from Vector Store
    chunk_ids = [chunk.embedding_id for chunk in document.chunks if chunk.embedding_id]
    if chunk_ids:
        await vector_store.delete(ids=chunk_ids)
        
    await db.delete(document)
    await db.commit()
    
    return {"status": "success"}

from pydantic import BaseModel

class MemoryCreate(BaseModel):
    title: str
    content: str
    tags: List[str] = []

class MemoryUpdate(BaseModel):
    title: str
    content: str
    tags: List[str] = []

@router.post("/memory", response_model=Any)
async def create_memory(
    memory_in: MemoryCreate,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    """
    Create a new memory (replaces old document-based memory creation).
    Respects user Auto-Approve settings.
    """
    from app.models.memory import Memory
    
    # Check Auto-Approve Setting
    auto_approve = True
    user_settings = current_user.settings
    
    if user_settings:
        if isinstance(user_settings, str):
            import json
            try:
                user_settings = json.loads(user_settings)
            except:
                user_settings = {}
                
        if isinstance(user_settings, dict):
            auto_approve = user_settings.get("auto_approve", True)
        
    initial_status = "approved" if auto_approve else "pending"
    # User upload: If approved, skip inbox. If pending, show in inbox.
    show_in_inbox = True if initial_status == "pending" else False
    
    # Rule: Anything external (identified by 'extension' tag) goes to inbox
    if memory_in.tags and "extension" in memory_in.tags:
        show_in_inbox = True
    
    embedding_id = str(uuid.uuid4())
    
    memory = Memory(
        title=memory_in.title,
        content=memory_in.content,
        user_id=current_user.id,
        tags=memory_in.tags,
        embedding_id=embedding_id,
        status=initial_status,
        show_in_inbox=show_in_inbox,
        source_llm="user-upload" # or extension
    )
    db.add(memory)
    await db.commit()
    await db.refresh(memory)

    # Trigger Background Analysis (Auto-Tagging + Similarity)
    background_tasks.add_task(run_metadata_extraction, memory.id, current_user.id)
    
    # Ingest only if approved
    if initial_status == "approved":
        ingest_memory_task.delay(
            memory_id=memory.id,
            user_id=current_user.id,
            content=memory_in.content,
            title=memory.title,
            tags=memory_in.tags,
            source="user-upload",
            doc_type="memory",
            mode="append"
        )

    # Trigger Dedupe Job (Background Celery)
    dedupe_memory_task.delay(memory.id)

    
    return {"status": "success", "document_id": memory.id, "chunks": 1 if initial_status == "approved" else 0}

@router.put("/{doc_id}", response_model=Any)
async def update_document(
    doc_id: int,
    memory: MemoryUpdate,
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    """
    Update a document (memory only). Re-chunks and updates vector store.
    """
    # Eager load chunks to avoid lazy loading issues in async
    from sqlalchemy.orm import selectinload
    result = await db.execute(
        select(Document)
        .options(selectinload(Document.chunks))
        .where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    document = result.scalars().first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Update document
    document.title = memory.title
    document.content = memory.content
    document.tags = memory.tags
    
    await db.commit()
    
    # Offload re-ingestion to background task
    from app.worker import ingest_memory_task
    ingest_memory_task.delay(
        memory_id=document.id,
        user_id=current_user.id,
        content=memory.content,
        title=document.title,
        tags=document.tags,
        source=document.source, # Keep original source or update?
        doc_type=document.doc_type, # file, memory, youtube
        mode="replace"
    )
    
    return {"status": "success", "document_id": document.id, "message": "Document updated and queued for re-processing"}


class SearchRequest(BaseModel):
    query: str
    top_k: int = 5

@router.post("/search", response_model=Any)
async def search_documents(
    request: SearchRequest,
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    """
    Semantic search over documents and memories.
    """
    # Use retrieval service for smart search
    # Use retrieval service for smart search
    results = await retrieval_service.search_memories(
        query=request.query,
        user_id=current_user.id,
        db=db,
        top_k=request.top_k
    )
    
    # Filter out non-serializable objects (like SQLAlchemy models)
    sanitized_results = []
    for res in results:
        sanitized_results.append({
            "text": res["text"],
            "score": res["score"],
            "metadata": res["metadata"]
        })
        
    return sanitized_results

from typing import List, Optional
class ChunkSchema(BaseModel):
    id: int
    text: str
    chunk_index: int
    document_id: Optional[int]
    memory_id: Optional[int]
    embedding_id: Optional[str]
    summary: Optional[str]
    generated_qas: Optional[Any] # Can be list or None
    entities: Optional[Any]
    trust_score: Optional[float]
    class Config:
        from_attributes = True

@router.get("/memory/{doc_id}/chunks", response_model=List[ChunkSchema])
async def get_document_chunks(
    doc_id: str,
    db: AsyncSession = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Any:
    """
    Get chunks for a specific memory/document (for Enrichment Sidebar).
    """
    # Parse ID and Build Query
    try:
        query = select(Chunk)
        
        if doc_id.startswith("mem_"):
            numeric_id = int(doc_id.split("_")[-1])
            query = query.where(Chunk.memory_id == numeric_id)
        elif doc_id.startswith("doc_"):
            numeric_id = int(doc_id.split("_")[-1])
            query = query.where(Chunk.document_id == numeric_id)
        else:
            # Fallback for legacy IDs (assume likely memory, or try both safest usage)
            # To be safe against collisions, maybe we should default to memory if unknown?
            # But legacy might be mixed. Sticking to OR for legacy integers only.
            numeric_id = int(doc_id)
            query = query.where(
                (Chunk.document_id == numeric_id) | (Chunk.memory_id == numeric_id)
            )
            
        query = query.order_by(Chunk.chunk_index)
            
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid Document ID format")
    
    result = await db.execute(query)
    chunks = result.scalars().all()
    
    return chunks
    
    if not results["documents"]:
        return []
        
    # Format results
    formatted_results = []
    # results["documents"] is a list of lists (one list per query)
    # results["metadatas"] is a list of lists
    
    docs = results["documents"][0]
    metas = results["metadatas"][0] if results["metadatas"] else [{}] * len(docs)
    
    for doc, meta in zip(docs, metas):
        formatted_results.append({
            "content": doc,
            "metadata": meta
        })
        
    return formatted_results


